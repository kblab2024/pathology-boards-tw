"""Convert a single handbook markdown file to .docx.

Lightweight, focused converter — handles only the constructs that appear
in this repo's handbook files: headings, paragraphs (with bold/em/inline
code), blockquotes, bullet lists, GFM tables, fenced code blocks, hr.

Not a general-purpose markdown-to-docx tool. Keep it that way.
"""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm
from markdown_it import MarkdownIt

from tools.render.config import HANDBOOK_DIR, RENDERED_DIR

CJK_FONT = "Noto Serif TC"
MONO_FONT = "Consolas"
ACCENT = RGBColor(0xB8, 0x23, 0x2F)        # 紅
HEADER_FILL = "B8232F"                       # 表頭背景
STRIPE_FILL = "F4EFE6"                       # 表格條紋
QUOTE_FILL = "F4EFE6"                        # blockquote 背景


def _set_run_font(run, name: str = CJK_FONT, size_pt: float | None = None,
                  bold: bool | None = None, color: RGBColor | None = None):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _shade_cell(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _shade_paragraph(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _set_cell_borders(cell, color: str = "C8C2B4", size: str = "4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcPr.append(tcBorders)
        tcBorders.append(b)


def _flatten_inline(inline_token) -> list[dict]:
    """Walk a markdown-it inline token's children and emit run-spec dicts.

    Each dict: {text, bold, italic, code}.
    """
    runs: list[dict] = []
    bold = False
    italic = False
    for tok in inline_token.children or []:
        t = tok.type
        if t == "text":
            runs.append({"text": tok.content, "bold": bold, "italic": italic, "code": False})
        elif t == "strong_open":
            bold = True
        elif t == "strong_close":
            bold = False
        elif t == "em_open":
            italic = True
        elif t == "em_close":
            italic = False
        elif t == "code_inline":
            runs.append({"text": tok.content, "bold": bold, "italic": italic, "code": True})
        elif t == "softbreak":
            runs.append({"text": " ", "bold": bold, "italic": italic, "code": False})
        elif t == "hardbreak":
            runs.append({"text": "\n", "bold": bold, "italic": italic, "code": False})
        # ignore link_open/close (no links expected in this content)
    return runs


def _add_runs(paragraph, runs: list[dict], base_size: float = 10.5,
              base_color: RGBColor | None = None):
    for spec in runs:
        text = spec["text"]
        if not text:
            continue
        run = paragraph.add_run(text)
        if spec["code"]:
            _set_run_font(run, name=MONO_FONT, size_pt=base_size,
                          bold=spec["bold"], color=base_color)
        else:
            _set_run_font(run, name=CJK_FONT, size_pt=base_size,
                          bold=spec["bold"], color=base_color)
        if spec["italic"]:
            run.italic = True


def _setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = CJK_FONT
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)
    rFonts.set(qn("w:ascii"), CJK_FONT)
    rFonts.set(qn("w:hAnsi"), CJK_FONT)
    return doc


def _add_heading(doc: Document, level: int, runs: list[dict]):
    sizes = {1: 22, 2: 16, 3: 13, 4: 11.5}
    size = sizes.get(level, 11)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    para.paragraph_format.space_after = Pt(6)
    for spec in runs:
        run = para.add_run(spec["text"])
        _set_run_font(run, name=CJK_FONT, size_pt=size, bold=True,
                      color=ACCENT if level <= 2 else None)
    if level == 1:
        # accent rule via bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "18")
        bottom.set(qn("w:color"), "C9A14A")
        pBdr.append(bottom)
        pPr.append(pBdr)


def _add_blockquote(doc: Document, runs: list[dict]):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    _shade_paragraph(para, QUOTE_FILL)
    _add_runs(para, runs, base_size=10, base_color=RGBColor(0x4A, 0x4A, 0x4A))


def _add_paragraph(doc: Document, runs: list[dict]):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    _add_runs(para, runs, base_size=10.5)


def _add_bullet(doc: Document, runs: list[dict]):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    _add_runs(para, runs, base_size=10.5)


def _add_code_block(doc: Document, content: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(8)
    _shade_paragraph(para, "F2F2F2")
    run = para.add_run(content.rstrip("\n"))
    _set_run_font(run, name=MONO_FONT, size_pt=9.5)


def _add_hr(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "C9A14A")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(doc: Document, header_rows: list[list[list[dict]]],
               body_rows: list[list[list[dict]]]):
    if not header_rows and not body_rows:
        return
    col_count = max(len(r) for r in (header_rows + body_rows))
    total_rows = len(header_rows) + len(body_rows)
    table = doc.add_table(rows=total_rows, cols=col_count)
    table.autofit = True
    table.style = "Table Grid"

    # header
    for r_i, row in enumerate(header_rows):
        for c_i, cell_runs in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _shade_cell(cell, HEADER_FILL)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            for spec in cell_runs:
                run = para.add_run(spec["text"])
                _set_run_font(run, name=CJK_FONT, size_pt=10, bold=True,
                              color=RGBColor(0xFF, 0xFF, 0xFF))

    # body
    for r_i, row in enumerate(body_rows):
        zebra = (r_i % 2 == 1)
        for c_i, cell_runs in enumerate(row):
            cell = table.rows[len(header_rows) + r_i].cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if zebra:
                _shade_cell(cell, STRIPE_FILL)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            _add_runs(para, cell_runs, base_size=10)

    # spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def convert(md_path: Path, docx_path: Path) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    md = MarkdownIt("commonmark", {"html": False}).enable("table").enable("strikethrough")
    tokens = md.parse(md_text)

    doc = _setup_document()

    i = 0
    in_blockquote = False
    in_list = False
    table_header: list[list[list[dict]]] = []
    table_body: list[list[list[dict]]] = []
    current_row: list[list[dict]] = []
    in_thead = False
    in_tbody = False

    while i < len(tokens):
        tok = tokens[i]
        t = tok.type

        if t == "heading_open":
            level = int(tok.tag[1])
            inline = tokens[i + 1]
            runs = _flatten_inline(inline)
            _add_heading(doc, level, runs)
            i += 3  # heading_open, inline, heading_close
            continue

        if t == "paragraph_open":
            inline = tokens[i + 1]
            runs = _flatten_inline(inline)
            if in_blockquote:
                _add_blockquote(doc, runs)
            elif in_list:
                _add_bullet(doc, runs)
            else:
                _add_paragraph(doc, runs)
            i += 3
            continue

        if t == "blockquote_open":
            in_blockquote = True
            i += 1
            continue
        if t == "blockquote_close":
            in_blockquote = False
            i += 1
            continue

        if t == "bullet_list_open":
            in_list = True
            i += 1
            continue
        if t == "bullet_list_close":
            in_list = False
            i += 1
            continue

        if t in ("list_item_open", "list_item_close"):
            i += 1
            continue

        if t == "fence" or t == "code_block":
            _add_code_block(doc, tok.content)
            i += 1
            continue

        if t == "hr":
            _add_hr(doc)
            i += 1
            continue

        if t == "table_open":
            table_header = []
            table_body = []
            i += 1
            continue
        if t == "table_close":
            _add_table(doc, table_header, table_body)
            table_header = []
            table_body = []
            i += 1
            continue
        if t == "thead_open":
            in_thead = True
            i += 1
            continue
        if t == "thead_close":
            in_thead = False
            i += 1
            continue
        if t == "tbody_open":
            in_tbody = True
            i += 1
            continue
        if t == "tbody_close":
            in_tbody = False
            i += 1
            continue
        if t == "tr_open":
            current_row = []
            i += 1
            continue
        if t == "tr_close":
            if in_thead:
                table_header.append(current_row)
            else:
                table_body.append(current_row)
            current_row = []
            i += 1
            continue
        if t in ("th_open", "td_open"):
            inline = tokens[i + 1]
            runs = _flatten_inline(inline)
            current_row.append(runs)
            i += 3  # *_open, inline, *_close
            continue

        # unhandled token: skip
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main():
    parser = argparse.ArgumentParser(description="Convert a handbook .md to .docx")
    parser.add_argument("--source", required=True,
                        help="Markdown filename inside handbook/ (e.g. ibd-summary.md)")
    parser.add_argument("--out", default=None,
                        help="Output .docx filename inside rendered/ (default: same stem)")
    args = parser.parse_args()

    md_path = HANDBOOK_DIR / args.source
    if not md_path.exists():
        raise SystemExit(f"source not found: {md_path}")

    out_name = args.out or (md_path.stem + ".docx")
    docx_path = RENDERED_DIR / out_name
    convert(md_path, docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    main()
